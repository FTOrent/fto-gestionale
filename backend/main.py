"""
FTO Gestionale Noleggio — FastAPI Backend v3.2
Tutti i campi del gestionale desktop, sicurezza completa.
"""
from fastapi import FastAPI, Depends, HTTPException, Request
from fastapi.responses import StreamingResponse
import io
from fastapi.security import OAuth2PasswordBearer
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from starlette.middleware.base import BaseHTTPMiddleware
from datetime import datetime, timedelta
from typing import Optional
import os, hashlib, platform, time
from collections import defaultdict
from jose import JWTError, jwt
from pydantic import BaseModel
import psycopg2, psycopg2.extras
from contextlib import contextmanager

SECRET_KEY         = os.getenv("SECRET_KEY", "fto-secret-2025")
ALGORITHM          = "HS256"
TOKEN_EXPIRE_HOURS = 12
DATABASE_URL       = os.getenv("DATABASE_URL", "")
SUPABASE_URL       = "https://nirauinyoiqtqcztgpjq.supabase.co"
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY", "")
STORAGE_BUCKET     = "documenti-noleggi"

ALLOWED_ORIGINS = [
    "https://web-production-eb644.up.railway.app",
    "https://gestionale.ftorent.it",
    "https://ftorent.it",
    "http://localhost:8000",
    "http://127.0.0.1:8000",
]

MAX_ATTEMPTS   = 5
LOCKOUT_SECS   = 900
RATE_LIMIT_RPM = 60

def _h(pw): return hashlib.sha256(pw.encode()).hexdigest()

USERS = {
    "ak":       {"hash": _h("ak47ven"),   "role": "admin",    "display": "Ak"},
    "toad":     {"hash": _h("caramello"), "role": "admin",    "display": "Toad"},
    "fuma":     {"hash": _h("Alex2001$"), "role": "admin",    "display": "Fuma"},
    "utente 1": {"hash": _h("ftorent$"),  "role": "employee", "display": "Utente 1"},
}

_login_attempts: dict = defaultdict(lambda: {"attempts": 0, "locked_until": 0.0})
_request_log:    dict = defaultdict(list)

def get_ip(request: Request) -> str:
    fwd = request.headers.get("x-forwarded-for")
    if fwd: return fwd.split(",")[0].strip()
    return request.client.host if request.client else "unknown"

def rate_limit(ip: str):
    now = time.time()
    _request_log[ip] = [t for t in _request_log[ip] if t > now - 60]
    if len(_request_log[ip]) >= RATE_LIMIT_RPM:
        raise HTTPException(429, "Troppe richieste.")
    _request_log[ip].append(now)

def check_brute(ip: str):
    s = _login_attempts[ip]
    if time.time() < s["locked_until"]:
        raise HTTPException(429, f"Bloccato per {int(s['locked_until']-time.time())}s.")

def fail_login(ip: str):
    s = _login_attempts[ip]; s["attempts"] += 1
    if s["attempts"] >= MAX_ATTEMPTS:
        s["locked_until"] = time.time() + LOCKOUT_SECS; s["attempts"] = 0

def ok_login(ip: str):
    _login_attempts[ip] = {"attempts": 0, "locked_until": 0.0}

class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request, call_next):
        r = await call_next(request)
        r.headers.update({
            "X-Content-Type-Options": "nosniff",
            "X-Frame-Options": "DENY",
            "X-XSS-Protection": "1; mode=block",
            "Referrer-Policy": "strict-origin-when-cross-origin",
            "Strict-Transport-Security": "max-age=31536000; includeSubDomains",
        })
        return r

@contextmanager
def db():
    conn = psycopg2.connect(DATABASE_URL, cursor_factory=psycopg2.extras.RealDictCursor)
    try:
        yield conn; conn.commit()
    except Exception:
        conn.rollback(); raise
    finally:
        conn.close()

def init_db():
    with db() as conn:
        conn.cursor().execute("""
        CREATE TABLE IF NOT EXISTS noleggi (
            id SERIAL PRIMARY KEY,
            -- contratto
            data_contratto TEXT DEFAULT '',
            -- cliente
            cl_nome TEXT DEFAULT '', cl_cognome TEXT DEFAULT '',
            cl_luogo TEXT DEFAULT '', cl_data_nascita TEXT DEFAULT '',
            cl_indirizzo TEXT DEFAULT '', cl_cf TEXT DEFAULT '',
            cl_doc TEXT DEFAULT '', cl_doc_da TEXT DEFAULT '',
            cl_doc_dt TEXT DEFAULT '', cl_pat TEXT DEFAULT '',
            cl_cat TEXT DEFAULT '', cl_pat_r TEXT DEFAULT '',
            cl_pat_s TEXT DEFAULT '',
            -- società
            soc_nome TEXT DEFAULT '', soc_sede TEXT DEFAULT '',
            soc_cf TEXT DEFAULT '', soc_piva TEXT DEFAULT '',
            -- guidatore
            gu_nome TEXT DEFAULT '', gu_cognome TEXT DEFAULT '',
            gu_luogo TEXT DEFAULT '', gu_data_nascita TEXT DEFAULT '',
            gu_indirizzo TEXT DEFAULT '', gu_cf TEXT DEFAULT '',
            gu_doc TEXT DEFAULT '', gu_doc_da TEXT DEFAULT '',
            gu_doc_dt TEXT DEFAULT '', gu_pat TEXT DEFAULT '',
            gu_cat TEXT DEFAULT '', gu_pat_r TEXT DEFAULT '',
            gu_pat_s TEXT DEFAULT '',
            -- conducente aggiuntivo
            ex_nome TEXT DEFAULT '', ex_cognome TEXT DEFAULT '',
            ex_luogo TEXT DEFAULT '', ex_data_nascita TEXT DEFAULT '',
            ex_indirizzo TEXT DEFAULT '', ex_cf TEXT DEFAULT '',
            ex_doc TEXT DEFAULT '', ex_doc_da TEXT DEFAULT '',
            ex_doc_dt TEXT DEFAULT '', ex_pat TEXT DEFAULT '',
            ex_cat TEXT DEFAULT '', ex_pat_r TEXT DEFAULT '',
            ex_pat_s TEXT DEFAULT '',
            -- auto
            marca_modello TEXT DEFAULT '', targa TEXT DEFAULT '',
            telaio TEXT DEFAULT '', cilindrata TEXT DEFAULT '',
            alimentazione TEXT DEFAULT '', anno TEXT DEFAULT '',
            colore TEXT DEFAULT '', km_iniziali TEXT DEFAULT '',
            -- date
            data_inizio TEXT DEFAULT '', ora_inizio TEXT DEFAULT '',
            data_fine TEXT DEFAULT '', ora_fine TEXT DEFAULT '',
            giorni TEXT DEFAULT '',
            -- tariffe
            prezzo_noleggio TEXT DEFAULT '', km_compresi TEXT DEFAULT '',
            deposito TEXT DEFAULT '', km_extra TEXT DEFAULT '',
            penale TEXT DEFAULT '', costo_violazione TEXT DEFAULT '',
            diff_carburante TEXT DEFAULT '', riconsegna_premium TEXT DEFAULT '',
            -- meta
            contratto_path TEXT DEFAULT '',
            verbale_stato TEXT DEFAULT '⏳ In attesa',
            verbale_path TEXT DEFAULT '',
            note TEXT DEFAULT '',
            nome_cognome TEXT GENERATED ALWAYS AS (cl_nome || ' ' || cl_cognome) STORED,
            created_at TIMESTAMPTZ DEFAULT NOW(),
            updated_at TIMESTAMPTZ DEFAULT NOW()
        );
        CREATE TABLE IF NOT EXISTS verbali (
            id SERIAL PRIMARY KEY, noleggio_id INTEGER,
            data TEXT, ora TEXT, targa TEXT,
            km_iniziali TEXT, km_attuali TEXT, carburante TEXT,
            tipo_riconsegna TEXT, pulizia TEXT, pulizia_premium TEXT,
            oggetti TEXT, oggetti_desc TEXT, danno TEXT, tipo_danno TEXT,
            posizione_danno TEXT, costo_danno TEXT, descrizione TEXT,
            operatore TEXT, stato TEXT,
            created_at TIMESTAMPTZ DEFAULT NOW()
        );
        CREATE TABLE IF NOT EXISTS auto (
            id SERIAL PRIMARY KEY, marca_modello TEXT, targa TEXT UNIQUE,
            telaio TEXT DEFAULT '', cilindrata TEXT DEFAULT '',
            alimentazione TEXT DEFAULT '', anno TEXT DEFAULT '',
            colore TEXT DEFAULT '', km INTEGER DEFAULT 0,
            updated_at TIMESTAMPTZ DEFAULT NOW()
        );
        CREATE TABLE IF NOT EXISTS access_log (
            id SERIAL PRIMARY KEY, ts TIMESTAMPTZ DEFAULT NOW(),
            username TEXT, role TEXT, success BOOLEAN,
            ip TEXT, hostname TEXT, os_info TEXT, machine TEXT
        );
        """)
        print("[DB] Tables ready")

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/api/auth/login")

def create_token(data: dict):
    return jwt.encode({**data, "exp": datetime.utcnow() + timedelta(hours=TOKEN_EXPIRE_HOURS)},
                      SECRET_KEY, algorithm=ALGORITHM)

def verify_token(token: str = Depends(oauth2_scheme)):
    try:
        p = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        if not p.get("sub"): raise HTTPException(401, "Token non valido")
        return p
    except JWTError:
        raise HTTPException(401, "Token scaduto")

def admin_only(p=Depends(verify_token)):
    if p.get("role") != "admin": raise HTTPException(403, "Solo admin")
    return p

app = FastAPI(title="FTO Gestionale", version="3.2", docs_url=None, redoc_url=None)
app.add_middleware(SecurityHeadersMiddleware)
app.add_middleware(CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_methods=["GET","POST","PUT","DELETE"],
    allow_headers=["Authorization","Content-Type"],
    allow_credentials=True)

@app.on_event("startup")
def startup():
    try: init_db()
    except Exception as e: print(f"[WARN] DB: {e}")

frontend_path = os.path.join(os.path.dirname(__file__), "..", "frontend")
static_path   = os.path.join(frontend_path, "static")
if os.path.exists(static_path):
    app.mount("/static", StaticFiles(directory=static_path), name="static")

@app.get("/")
def root():
    idx = os.path.join(frontend_path, "templates", "index.html")
    if os.path.exists(idx):
        return FileResponse(idx, headers={
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0"
        })
    return {"status": "FTO v3.2"}

# ── AUTH ──────────────────────────────────────────────────────
class LoginReq(BaseModel):
    username: str; password: str
    ip: Optional[str] = ""; hostname: Optional[str] = ""

@app.post("/api/auth/login")
def login(req: LoginReq, request: Request):
    ip = get_ip(request)
    check_brute(ip)
    username = req.username.strip().lower()
    user = USERS.get(username)
    success = user is not None and user["hash"] == _h(req.password)
    try:
        with db() as conn:
            conn.cursor().execute(
                "INSERT INTO access_log(username,role,success,ip,hostname,os_info,machine)"
                " VALUES(%s,%s,%s,%s,%s,%s,%s)",
                (username, user["role"] if user else "?", success,
                 ip, req.hostname, platform.system(), platform.machine()))
    except: pass
    if not success:
        fail_login(ip)
        raise HTTPException(401, "Username o password errati")
    ok_login(ip)
    return {"access_token": create_token({"sub":username,"role":user["role"],"display":user["display"]}),
            "token_type":"bearer","display":user["display"],"role":user["role"]}

@app.get("/api/auth/me")
def me(p=Depends(verify_token)):
    return {"username":p["sub"],"display":p["display"],"role":p["role"]}

# ── NOLEGGI ───────────────────────────────────────────────────
class NoleggioIn(BaseModel):
    data_contratto: Optional[str]=""
    cl_nome: Optional[str]=""; cl_cognome: Optional[str]=""
    cl_luogo: Optional[str]=""; cl_data_nascita: Optional[str]=""
    cl_indirizzo: Optional[str]=""; cl_cf: Optional[str]=""
    cl_doc: Optional[str]=""; cl_doc_da: Optional[str]=""
    cl_doc_dt: Optional[str]=""; cl_pat: Optional[str]=""
    cl_cat: Optional[str]=""; cl_pat_r: Optional[str]=""
    cl_pat_s: Optional[str]=""
    soc_nome: Optional[str]=""; soc_sede: Optional[str]=""
    soc_cf: Optional[str]=""; soc_piva: Optional[str]=""
    gu_nome: Optional[str]=""; gu_cognome: Optional[str]=""
    gu_luogo: Optional[str]=""; gu_data_nascita: Optional[str]=""
    gu_indirizzo: Optional[str]=""; gu_cf: Optional[str]=""
    gu_doc: Optional[str]=""; gu_doc_da: Optional[str]=""
    gu_doc_dt: Optional[str]=""; gu_pat: Optional[str]=""
    gu_cat: Optional[str]=""; gu_pat_r: Optional[str]=""
    gu_pat_s: Optional[str]=""
    ex_nome: Optional[str]=""; ex_cognome: Optional[str]=""
    ex_luogo: Optional[str]=""; ex_data_nascita: Optional[str]=""
    ex_indirizzo: Optional[str]=""; ex_cf: Optional[str]=""
    ex_doc: Optional[str]=""; ex_doc_da: Optional[str]=""
    ex_doc_dt: Optional[str]=""; ex_pat: Optional[str]=""
    ex_cat: Optional[str]=""; ex_pat_r: Optional[str]=""
    ex_pat_s: Optional[str]=""
    marca_modello: Optional[str]=""; targa: Optional[str]=""
    telaio: Optional[str]=""; cilindrata: Optional[str]=""
    alimentazione: Optional[str]=""; anno: Optional[str]=""
    colore: Optional[str]=""; km_iniziali: Optional[str]=""
    data_inizio: Optional[str]=""; ora_inizio: Optional[str]=""
    data_fine: Optional[str]=""; ora_fine: Optional[str]=""
    giorni: Optional[str]=""
    prezzo_noleggio: Optional[str]=""; km_compresi: Optional[str]=""
    deposito: Optional[str]=""; km_extra: Optional[str]=""
    penale: Optional[str]=""; costo_violazione: Optional[str]=""
    diff_carburante: Optional[str]=""; riconsegna_premium: Optional[str]=""
    contratto_path: Optional[str]=""; note: Optional[str]=""

NOLEGGIO_COLS = [
    "data_contratto",
    "cl_nome","cl_cognome","cl_luogo","cl_data_nascita","cl_indirizzo","cl_cf",
    "cl_doc","cl_doc_da","cl_doc_dt","cl_pat","cl_cat","cl_pat_r","cl_pat_s",
    "soc_nome","soc_sede","soc_cf","soc_piva",
    "gu_nome","gu_cognome","gu_luogo","gu_data_nascita","gu_indirizzo","gu_cf",
    "gu_doc","gu_doc_da","gu_doc_dt","gu_pat","gu_cat","gu_pat_r","gu_pat_s",
    "ex_nome","ex_cognome","ex_luogo","ex_data_nascita","ex_indirizzo","ex_cf",
    "ex_doc","ex_doc_da","ex_doc_dt","ex_pat","ex_cat","ex_pat_r","ex_pat_s",
    "marca_modello","targa","telaio","cilindrata","alimentazione","anno","colore","km_iniziali",
    "data_inizio","ora_inizio","data_fine","ora_fine","giorni",
    "prezzo_noleggio","km_compresi","deposito","km_extra","penale",
    "costo_violazione","diff_carburante","riconsegna_premium",
    "contratto_path","note",
]

@app.get("/api/noleggi")
def list_noleggi(request: Request, q:Optional[str]=None, nome:Optional[str]=None,
    auto:Optional[str]=None, da:Optional[str]=None, a:Optional[str]=None,
    prezzo_min:Optional[float]=None, prezzo_max:Optional[float]=None,
    stato:Optional[str]=None, p=Depends(verify_token)):
    rate_limit(get_ip(request))
    with db() as conn:
        cur = conn.cursor()
        sql = "SELECT * FROM noleggi WHERE 1=1"; params = []
        if q:
            sql += " AND (nome_cognome ILIKE %s OR marca_modello ILIKE %s OR targa ILIKE %s)"
            params += [f"%{q}%"]*3
        if nome: sql += " AND nome_cognome ILIKE %s"; params.append(f"%{nome}%")
        if auto:
            sql += " AND (marca_modello ILIKE %s OR targa ILIKE %s)"
            params += [f"%{auto}%"]*2
        if da:
            sql += " AND TO_DATE(NULLIF(data_inizio,''),'DD/MM/YYYY') >= TO_DATE(%s,'DD/MM/YYYY')"
            params.append(da)
        if a:
            sql += " AND TO_DATE(NULLIF(data_inizio,''),'DD/MM/YYYY') <= TO_DATE(%s,'DD/MM/YYYY')"
            params.append(a)
        if prezzo_min is not None:
            sql += " AND NULLIF(prezzo_noleggio,'')::numeric >= %s"; params.append(prezzo_min)
        if prezzo_max is not None:
            sql += " AND NULLIF(prezzo_noleggio,'')::numeric <= %s"; params.append(prezzo_max)
        if stato and stato != "tutti":
            if stato == "attesa": sql += " AND (verbale_stato ILIKE '%attesa%' OR verbale_stato='')"
            elif stato == "ok":   sql += " AND verbale_stato ILIKE '%OK%'"
            elif stato == "danno": sql += " AND verbale_stato ILIKE '%Danno%'"
        cur.execute(sql + " ORDER BY id DESC", params)
        return cur.fetchall()

@app.get("/api/noleggi/{id}")
def get_noleggio(id:int, request:Request, p=Depends(verify_token)):
    rate_limit(get_ip(request))
    with db() as conn:
        cur = conn.cursor(); cur.execute("SELECT * FROM noleggi WHERE id=%s",(id,))
        row = cur.fetchone()
        if not row: raise HTTPException(404,"Non trovato")
        return row

@app.post("/api/noleggi", status_code=201)
def create_noleggio(data: NoleggioIn, request:Request, p=Depends(verify_token)):
    rate_limit(get_ip(request))
    vals = [getattr(data, c) for c in NOLEGGIO_COLS]
    cols_sql = ",".join(NOLEGGIO_COLS)
    ph = ",".join(["%s"]*len(NOLEGGIO_COLS))
    with db() as conn:
        cur = conn.cursor()
        cur.execute(f"INSERT INTO noleggi ({cols_sql}) VALUES ({ph}) RETURNING id", vals)
        return {"id": cur.fetchone()["id"]}

@app.put("/api/noleggi/{id}")
def update_noleggio(id:int, data:NoleggioIn, request:Request, p=Depends(admin_only)):
    rate_limit(get_ip(request))
    update_cols = [c for c in NOLEGGIO_COLS if c != "data_contratto"]
    set_sql = ",".join(f"{c}=%s" for c in update_cols) + ",updated_at=NOW()"
    vals = [getattr(data, c) for c in update_cols] + [id]
    with db() as conn:
        conn.cursor().execute(f"UPDATE noleggi SET {set_sql} WHERE id=%s", vals)
        return {"ok": True}

@app.delete("/api/noleggi/{id}")
def delete_noleggio(id:int, request:Request, p=Depends(admin_only)):
    rate_limit(get_ip(request))
    with db() as conn:
        conn.cursor().execute("DELETE FROM noleggi WHERE id=%s",(id,))
        return {"ok": True}

# ── VERBALI ───────────────────────────────────────────────────
class VerbaleIn(BaseModel):
    noleggio_id:int; data:str; ora:str; targa:str
    km_iniziali:str; km_attuali:str; carburante:str
    tipo_riconsegna:str; pulizia:str; pulizia_premium:str
    oggetti:str; oggetti_desc:Optional[str]=""
    danno:str; tipo_danno:Optional[str]=""
    posizione_danno:Optional[str]=""; costo_danno:Optional[str]=""
    descrizione:Optional[str]=""; operatore:str; stato:str

@app.post("/api/verbali", status_code=201)
def create_verbale(data:VerbaleIn, request:Request, p=Depends(verify_token)):
    rate_limit(get_ip(request))
    with db() as conn:
        cur = conn.cursor()
        cur.execute("""INSERT INTO verbali
          (noleggio_id,data,ora,targa,km_iniziali,km_attuali,carburante,
           tipo_riconsegna,pulizia,pulizia_premium,oggetti,oggetti_desc,
           danno,tipo_danno,posizione_danno,costo_danno,descrizione,operatore,stato)
          VALUES(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) RETURNING id""",
          (data.noleggio_id,data.data,data.ora,data.targa,data.km_iniziali,
           data.km_attuali,data.carburante,data.tipo_riconsegna,data.pulizia,
           data.pulizia_premium,data.oggetti,data.oggetti_desc,data.danno,
           data.tipo_danno,data.posizione_danno,data.costo_danno,
           data.descrizione,data.operatore,data.stato))
        vid = cur.fetchone()["id"]
        cur.execute("UPDATE noleggi SET verbale_stato=%s,updated_at=NOW() WHERE id=%s",
                    (data.stato,data.noleggio_id))
        try:
            km_new = int(float(data.km_attuali.replace(",","")))
            cur.execute("UPDATE auto SET km=%s,updated_at=NOW() WHERE targa=%s",(km_new,data.targa))
        except: pass
        return {"id":vid}

@app.get("/api/verbali/pending")
def pending_verbali(request:Request, p=Depends(verify_token)):
    rate_limit(get_ip(request))
    with db() as conn:
        cur = conn.cursor()
        cur.execute("""SELECT id,nome_cognome,marca_modello,targa,data_inizio,km_iniziali
          FROM noleggi WHERE verbale_stato ILIKE '%attesa%' OR verbale_stato=''
          ORDER BY id DESC""")
        return cur.fetchall()

# ── AUTO ──────────────────────────────────────────────────────
class AutoIn(BaseModel):
    marca_modello:str; targa:str
    telaio:Optional[str]=""; cilindrata:Optional[str]=""
    alimentazione:Optional[str]=""; anno:Optional[str]=""
    colore:Optional[str]=""; km:Optional[int]=0

@app.get("/api/auto")
def list_auto(request:Request, p=Depends(verify_token)):
    rate_limit(get_ip(request))
    with db() as conn:
        cur = conn.cursor(); cur.execute("SELECT * FROM auto ORDER BY marca_modello")
        return cur.fetchall()

@app.post("/api/auto", status_code=201)
def create_auto(data:AutoIn, request:Request, p=Depends(admin_only)):
    rate_limit(get_ip(request))
    with db() as conn:
        cur = conn.cursor()
        cur.execute("""INSERT INTO auto(marca_modello,targa,telaio,cilindrata,alimentazione,anno,colore,km)
          VALUES(%s,%s,%s,%s,%s,%s,%s,%s)
          ON CONFLICT(targa) DO UPDATE SET marca_modello=EXCLUDED.marca_modello,
          km=EXCLUDED.km,updated_at=NOW() RETURNING id""",
          (data.marca_modello,data.targa,data.telaio,data.cilindrata,
           data.alimentazione,data.anno,data.colore,data.km))
        auto_id = cur.fetchone()["id"]
        log_change(p["sub"], p["role"], "CREA", "auto", auto_id,
            f"Aggiunto veicolo {data.marca_modello} ({data.targa})",
            get_ip(request))
        return {"id": auto_id}

# ── LOGS ──────────────────────────────────────────────────────
@app.get("/api/logs")
def get_logs(request:Request, p=Depends(admin_only)):
    rate_limit(get_ip(request))
    with db() as conn:
        cur = conn.cursor()
        cur.execute("""
            SELECT id,
                   TO_CHAR(ts AT TIME ZONE 'Europe/Rome', 'DD/MM/YYYY HH24:MI:SS') AS ts,
                   username, role, success, ip, hostname, os_info, machine
            FROM access_log
            ORDER BY id DESC LIMIT 200
        """)
        rows = cur.fetchall()
        return [dict(r) for r in rows]


# ── CONTRATTO PDF ─────────────────────────────────────────────
@app.get("/api/contratto/{id}")
def genera_contratto(id: int, request: Request):
    """Generate contract DOCX from template, filling all {{placeholders}}"""
    # Auth via header
    from fastapi.security.utils import get_authorization_scheme_param
    auth = request.headers.get("Authorization", "")
    _, tok = get_authorization_scheme_param(auth)
    try:
        p = verify_token(tok) if tok else None
    except:
        raise HTTPException(401, "Non autorizzato")
    if not p:
        raise HTTPException(401, "Non autorizzato")

    rate_limit(get_ip(request))
    with db() as conn:
        cur = conn.cursor()
        cur.execute("SELECT * FROM noleggi WHERE id=%s", (id,))
        r = cur.fetchone()
        if not r:
            raise HTTPException(404, "Noleggio non trovato")

    r = dict(r)

    def v(key, default="—"):
        val = r.get(key, "") or ""
        if str(val).strip() in ("", "nan", "None"):
            return default
        return str(val).strip()

    # Map all template placeholders to DB fields
    ctx = {
        "cliente_nome_cognome":                          f"{v('cl_nome')} {v('cl_cognome')}".strip(),
        "cliente_luogo_nascita":                         v("cl_luogo"),
        "cliente_data_nascita":                          v("cl_data_nascita"),
        "cliente_indirizzo_residenza":                   v("cl_indirizzo"),
        "cliente_codice_fiscale":                        v("cl_cf"),
        "cliente_numero_documento":                      v("cl_doc"),
        "cliente_documento_rilasciato_da":               v("cl_doc_da"),
        "cliente_documento_data_rilascio":               v("cl_doc_dt"),
        "cliente_numero_patente":                        v("cl_pat"),
        "cliente_categoria_patente":                     v("cl_cat"),
        "cliente_patente_data_rilascio":                 v("cl_pat_r"),
        "cliente_patente_scadenza":                      v("cl_pat_s"),
        "societa_nome":                                  v("soc_nome"),
        "societa_sede":                                  v("soc_sede"),
        "societa_codice_fiscale":                        v("soc_cf"),
        "societa_partita_iva":                           v("soc_piva"),
        "guidatore_nome_cognome":                        f"{v('gu_nome')} {v('gu_cognome')}".strip(),
        "guidatore_luogo_nascita":                       v("gu_luogo"),
        "guidatore_data_nascita":                        v("gu_data_nascita"),
        "guidatore_indirizzo_residenza":                 v("gu_indirizzo"),
        "guidatore_codice_fiscale":                      v("gu_cf"),
        "guidatore_numero_documento":                    v("gu_doc"),
        "guidatore_documento_rilasciato_da":             v("gu_doc_da"),
        "guidatore_documento_data_rilascio":             v("gu_doc_dt"),
        "guidatore_numero_patente":                      v("gu_pat"),
        "guidatore_categoria_patente":                   v("gu_cat"),
        "guidatore_patente_data_rilascio":               v("gu_pat_r"),
        "guidatore_patente_scadenza":                    v("gu_pat_s"),
        "conducente_aggiuntivo_nome_cognome":            f"{v('ex_nome')} {v('ex_cognome')}".strip(),
        "conducente_aggiuntivo_luogo_nascita":           v("ex_luogo"),
        "conducente_aggiuntivo_data_nascita":            v("ex_data_nascita"),
        "conducente_aggiuntivo_indirizzo_residenza":     v("ex_indirizzo"),
        "conducente_aggiuntivo_codice_fiscale":          v("ex_cf"),
        "conducente_aggiuntivo_numero_documento":        v("ex_doc"),
        "conducente_aggiuntivo_documento_rilasciato_da": v("ex_doc_da"),
        "conducente_aggiuntivo_documento_data_rilascio": v("ex_doc_dt"),
        "conducente_aggiuntivo_numero_patente":          v("ex_pat"),
        "conducente_aggiuntivo_categoria_patente":       v("ex_cat"),
        "conducente_aggiuntivo_patente_data_rilascio":   v("ex_pat_r"),
        "conducente_aggiuntivo_patente_scadenza":        v("ex_pat_s"),
        "marca_modello":        v("marca_modello"),
        "targa":                v("targa"),
        "telaio":               v("telaio"),
        "cilindrata":           v("cilindrata"),
        "alimentazione":        v("alimentazione"),
        "anno":                 v("anno"),
        "colore":               v("colore"),
        "km":                   v("km_iniziali"),
        "data_inizio":          v("data_inizio"),
        "ora_inizio":           v("ora_inizio"),
        "data_fine":            v("data_fine"),
        "ora_fine":             v("ora_fine"),
        "prezzo_noleggio":      v("prezzo_noleggio"),
        "km_compresi":          v("km_compresi"),
        "km_extra":             v("km_extra"),
        "riconsegna_premium":   v("riconsegna_premium"),
        "penale":               v("penale"),
        "deposito":             v("deposito"),
        "costo_violazione":     v("costo_violazione"),
        "diff_carburante":      v("diff_carburante"),
        "firma_data":           v("data_contratto") if v("data_contratto") != "—" else datetime.today().strftime("%d/%m/%Y"),
    }

    # Load the Word template
    import os
    from docx import Document as DocxDoc
    # Try multiple possible paths
    base = os.path.dirname(os.path.abspath(__file__))
    possible = [
        os.path.join(base, "templates", "contratto_template.docx"),
        os.path.join(base, "template", "contratto_template.docx"),
        "/app/backend/templates/contratto_template.docx",
        "/app/backend/template/contratto_template.docx",
    ]
    template_path = None
    for p in possible:
        if os.path.exists(p):
            template_path = p
            break
    if not template_path:
        searched = ", ".join(possible)
        raise HTTPException(500, f"Template non trovato. Cercato in: {searched}")

    doc = DocxDoc(template_path)

    def fill_paragraph(para):
        """Replace {{key}} in paragraph, preserving formatting."""
        for key, val in ctx.items():
            ph = "{{" + key + "}}"
            if ph not in para.text:
                continue
            # Rebuild runs preserving format of first run
            full = "".join(run.text for run in para.runs)
            if ph in full:
                new_full = full.replace(ph, val)
                # Put everything in first run, clear the rest
                if para.runs:
                    para.runs[0].text = new_full
                    for run in para.runs[1:]:
                        run.text = ""

    for para in doc.paragraphs:
        fill_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    fill_paragraph(para)

    # Return as .docx download
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    cl = f"{r.get('cl_nome','')}_{r.get('cl_cognome','')}".replace(" ", "_")
    nome_file = f"contratto_{cl}_{id}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={nome_file}"}
    )


@app.get("/api/change-log")
def get_change_log(request: Request, p=Depends(admin_only)):
    rate_limit(get_ip(request))
    with db() as conn:
        cur = conn.cursor()
        cur.execute("""
            SELECT id,
                   TO_CHAR(ts AT TIME ZONE 'Europe/Rome', 'DD/MM/YYYY HH24:MI:SS') AS ts,
                   username, role, action, entity, entity_id, description, ip
            FROM change_log
            ORDER BY ts DESC
            LIMIT 200
        """)
        return cur.fetchall()

# ── SUPABASE STORAGE HELPERS ──────────────────────────────────
import httpx

def log_change(username, role, action, entity, entity_id, description, ip):
    """Log a change action to change_log table."""
    try:
        with db() as conn:
            cur = conn.cursor()
            cur.execute("""INSERT INTO change_log(username,role,action,entity,entity_id,description,ip)
                VALUES(%s,%s,%s,%s,%s,%s,%s)""",
                (username, role, action, entity, entity_id, description, ip))
    except Exception as e:
        print(f"[LOG] Error logging change: {e}")

def storage_upload(path: str, data: bytes, content_type: str) -> bool:
    """Upload file to Supabase Storage."""
    url = f"{SUPABASE_URL}/storage/v1/object/{STORAGE_BUCKET}/{path}"
    headers = {
        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
        "Content-Type": content_type,
        "x-upsert": "true",
    }
    r = httpx.put(url, content=data, headers=headers, timeout=30)
    return r.status_code in (200, 201)

def storage_get_url(path: str) -> str:
    """Get signed URL for private file (valid 1 hour)."""
    url = f"{SUPABASE_URL}/storage/v1/object/sign/{STORAGE_BUCKET}/{path}"
    headers = {"Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}
    r = httpx.post(url, json={"expiresIn": 3600}, headers=headers, timeout=10)
    if r.status_code == 200:
        return SUPABASE_URL + "/storage/v1" + r.json().get("signedURL", "")
    return ""

def storage_list(prefix: str) -> list:
    """List files in a folder."""
    url = f"{SUPABASE_URL}/storage/v1/object/list/{STORAGE_BUCKET}"
    headers = {"Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"}
    r = httpx.post(url, json={"prefix": prefix, "limit": 100}, headers=headers, timeout=10)
    if r.status_code == 200:
        return r.json()
    return []

# ── UPLOAD FOTO ───────────────────────────────────────────────
from fastapi import UploadFile, File, Form
from typing import List

@app.post("/api/noleggi/{id}/foto")
async def upload_foto(
    id: int,
    file: UploadFile = File(...),
    tipo: str = Form(...),   # es: "cl_ci_f", "cl_pa_r", "gu_pp_f" ecc
    request: Request = None,
    p=Depends(verify_token)
):
    rate_limit(get_ip(request))
    # Verify noleggio exists
    with db() as conn:
        cur = conn.cursor()
        cur.execute("SELECT id FROM noleggi WHERE id=%s", (id,))
        if not cur.fetchone():
            raise HTTPException(404, "Noleggio non trovato")

    data = await file.read()
    ext  = file.filename.split(".")[-1].lower() if "." in file.filename else "jpg"
    path = f"noleggio_{id}/{tipo}.{ext}"
    ct   = file.content_type or "image/jpeg"

    ok = storage_upload(path, data, ct)
    if not ok:
        raise HTTPException(500, "Errore upload foto")
    return {"ok": True, "path": path}

# ── GET FOTO LIST ─────────────────────────────────────────────
@app.get("/api/noleggi/{id}/foto")
def get_foto(id: int, request: Request, p=Depends(verify_token)):
    rate_limit(get_ip(request))
    files = storage_list(f"noleggio_{id}/")
    result = []
    for f in files:
        name = f.get("name","")
        path = f"noleggio_{id}/{name}"
        url  = storage_get_url(path)
        if url:
            result.append({"name": name, "url": url, "path": path})
    return result

# ── GET SINGLE FOTO ───────────────────────────────────────────
@app.get("/api/foto")
def get_foto_url(path: str, request: Request, p=Depends(verify_token)):
    rate_limit(get_ip(request))
    url = storage_get_url(path)
    if not url:
        raise HTTPException(404, "Foto non trovata")
    from fastapi.responses import RedirectResponse
    return RedirectResponse(url)
